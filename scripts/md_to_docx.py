"""
md_to_docx.py  –  Convert corpus/spec/FUNCTIONAL_SPEC.md → corpus/spec/FUNCTIONAL_SPEC.docx

Usage:
    python scripts/md_to_docx.py

Requires: python-docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SPEC_MD   = Path("corpus/spec/FUNCTIONAL_SPEC.md")
SPEC_DOCX = Path("corpus/spec/FUNCTIONAL_SPEC.docx")

SERIF_FONT = "Georgia"
BODY_PT    = 11
RULE_RE    = re.compile(r"^\*\*BR-\d{3}\.\*\*")

# ── helpers ──────────────────────────────────────────────────────────────────

def _set_font(run, name: str, size_pt: int, bold: bool = False,
              color: RGBColor | None = None):
    run.font.name       = name
    run.font.size       = Pt(size_pt)
    run.font.bold       = bold
    if color:
        run.font.color.rgb = color

def _para_font(paragraph, name: str, size_pt: int):
    """Apply font to every run in paragraph (for cases where we add runs manually)."""
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size_pt)


def _add_run_rich(para, text: str, serif: str, size_pt: int):
    """
    Parse inline markdown bold (**...**) and italic (*...*) inside *text*
    and add styled runs to *para*.
    """
    # tokenise: split on **...** or *...*
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            run = para.add_run(tok[2:-2])
            _set_font(run, serif, size_pt, bold=True)
        elif tok.startswith("*") and tok.endswith("*"):
            run = para.add_run(tok[1:-1])
            run.font.name  = serif
            run.font.size  = Pt(size_pt)
            run.font.italic = True
        else:
            run = para.add_run(tok)
            _set_font(run, serif, size_pt)


def _set_para_spacing(para, before: int = 0, after: int = 6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _add_title_page(doc: Document):
    """Insert a corporate-style title page."""
    doc.add_paragraph()  # top margin breathing room

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Company System")
    run.font.name  = "Calibri"
    run.font.size  = Pt(28)
    run.font.bold  = True

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run("Functional Specification")
    run2.font.name = "Calibri"
    run2.font.size = Pt(20)
    run2.font.bold = True

    doc.add_paragraph()  # spacer

    meta = [
        ("Document Reference", "COSYS-FS-001"),
        ("Version",            "1.0"),
        ("Date",               "1998"),
        ("Status",             "Baseline"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(f"{label}: ")
        r_label.font.name = "Calibri"
        r_label.font.size = Pt(13)
        r_label.font.bold = True
        r_val = p.add_run(value)
        r_val.font.name = "Calibri"
        r_val.font.size = Pt(13)

    # page break after title page
    doc.add_page_break()


def _add_table_from_md(doc: Document, header_row: list[str], rows: list[list[str]]):
    """Render a markdown pipe table as a Word table."""
    ncols = len(header_row)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"

    # header
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header_row):
        hdr_cells[i].text = txt
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.name = SERIF_FONT
            run.font.size = Pt(10)
            run.font.bold = True

    # body rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, txt in enumerate(row):
            row_cells[c_idx].text = txt
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.name = SERIF_FONT
                run.font.size = Pt(10)

    doc.add_paragraph()  # space after table


# ── main converter ────────────────────────────────────────────────────────────

def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc   = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    _add_title_page(doc)

    # State for table accumulation
    in_table      = False
    table_headers: list[str] = []
    table_rows:    list[list[str]] = []

    def flush_table():
        nonlocal in_table, table_headers, table_rows
        if in_table and table_headers:
            _add_table_from_md(doc, table_headers, table_rows)
        in_table      = False
        table_headers = []
        table_rows    = []

    def parse_pipe_row(line: str) -> list[str]:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── blank line ──
        if not line.strip():
            if in_table:
                flush_table()
            i += 1
            continue

        # ── horizontal rule ──
        if re.match(r"^-{3,}$", line.strip()):
            if in_table:
                flush_table()
            i += 1
            continue

        # ── heading ──
        m_hd = re.match(r"^(#{1,6})\s+(.*)", line)
        if m_hd:
            if in_table:
                flush_table()
            level = len(m_hd.group(1))
            text  = m_hd.group(2).strip()
            # map to Word heading styles
            style_name = f"Heading {level}"
            try:
                para = doc.add_paragraph(style=style_name)
            except KeyError:
                para = doc.add_paragraph(style="Heading 3")
            para.clear()
            # strip markdown bold from heading text
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            run   = para.add_run(clean)
            # heading font sizes
            sizes = {1: 18, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}
            run.font.name = "Calibri"
            run.font.size = Pt(sizes.get(level, 12))
            run.font.bold = True
            i += 1
            continue

        # ── table row ──
        if line.strip().startswith("|"):
            cells = parse_pipe_row(line)
            if not in_table:
                # first row → header
                in_table      = True
                table_headers = cells
                table_rows    = []
                # skip separator row (next line)
                if i + 1 < len(lines) and re.match(r"^\|[-| ]+\|$", lines[i + 1].strip()):
                    i += 2
                else:
                    i += 1
            else:
                table_rows.append(cells)
                i += 1
            continue

        # ── if we were in a table and hit a non-table line, flush ──
        if in_table:
            flush_table()

        # ── rule paragraph (BR-nnn) ──
        if RULE_RE.match(line):
            para = doc.add_paragraph()
            _set_para_spacing(para, before=4, after=4)
            # Strip outer ** from the BR-nnn prefix
            # e.g. "**BR-001.**  rest of text"
            m_rule = re.match(r"\*\*(BR-\d{3}\.)\*\*\s*(.*)", line)
            if m_rule:
                prefix = m_rule.group(1)
                rest   = m_rule.group(2)
                r_pre  = para.add_run(prefix + "  ")
                _set_font(r_pre, SERIF_FONT, BODY_PT, bold=True)
                _add_run_rich(para, rest, SERIF_FONT, BODY_PT)
            else:
                _add_run_rich(para, line, SERIF_FONT, BODY_PT)
            i += 1
            continue

        # ── list item ──
        m_li = re.match(r"^[-*]\s+(.*)", line)
        if m_li:
            para = doc.add_paragraph(style="List Bullet")
            para.clear()
            _set_para_spacing(para, before=0, after=2)
            _add_run_rich(para, m_li.group(1), SERIF_FONT, BODY_PT)
            i += 1
            continue

        # ── italic-only line (e.g. *End of spec*) ──
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            para = doc.add_paragraph()
            _set_para_spacing(para, before=12, after=0)
            run  = para.add_run(line.strip()[1:-1])
            run.font.name   = SERIF_FONT
            run.font.size   = Pt(BODY_PT)
            run.font.italic = True
            i += 1
            continue

        # ── plain body paragraph ──
        para = doc.add_paragraph()
        _set_para_spacing(para, before=0, after=6)
        _add_run_rich(para, line, SERIF_FONT, BODY_PT)
        i += 1

    if in_table:
        flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Saved: {docx_path}  ({docx_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    convert(SPEC_MD, SPEC_DOCX)
